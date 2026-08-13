import sys
import docx
import io
import re
from PIL import Image
import numpy as np
from typing import Dict, List, Tuple
import easyocr

sys.stdout.reconfigure(encoding='utf-8')

# Initialize EasyOCR reader once
ocr_reader = easyocr.Reader(['en'], gpu=False, verbose=False)

# Ground-truth annotated benchmark dataset from Red Herring Prospectus.docx
# Spans paragraphs, tables, and image card OCR text across all required PII categories.
BENCHMARK_GROUND_TRUTH = [
    # Category, Entity Value, Description
    ("PERSON", "Sarthak Malvadkar", "Contact Person"),
    ("PERSON", "VISHAL SINGH", "PAN Card Name"),
    ("PERSON", "SUGRIV SINGH", "PAN Card Father Name"),
    ("PERSON", "MERAJ KHAN", "Aadhaar Card Name"),
    ("PERSON", "Sudhdan Khan", "Aadhaar Card Father Name"),

    ("EMAIL_ADDRESS", "cs.connect@kshinternational.com", "Company Email"),
    ("EMAIL_ADDRESS", "tininfo@nsdl.co.in", "PAN Services Email"),

    ("PHONE_NUMBER", "+ 91 20 4505 3237", "Company Telephone"),
    ("PHONE_NUMBER", "91-20-2721 8080", "PAN Unit Telephone"),
    ("PHONE_NUMBER", "91-20-2721 8081", "PAN Unit Fax"),

    ("INDIAN_PAN", "NBWPS19SIN", "PAN Number"),
    ("INDIAN_AADHAAR", "2943 6593 3461", "Aadhaar Number"),

    ("ADDRESS", "11/3, 11/4 and 11/5, Village Birdewadi, Chakan Taluka - Khed, Pune – 410 501", "Registered Address"),
    ("ADDRESS", "201, Tower 2, Montreal Business Centre, Off Pallod Farms, Baner, Pune – 411 045", "Corporate Address"),
    ("ADDRESS", "saray dan shah KATRAULI", "Aadhaar Address"),

    ("DATE_OF_BIRTH", "06/05/2000", "PAN Card DOB"),
    ("DATE_OF_BIRTH", "12/12/1988", "Aadhaar Card DOB"),
    ("DATE_TIME", "July 30, 1979", "Incorporation Date"),
    ("DATE_TIME", "June 1, 1996", "Resolution Date"),
    ("DATE_TIME", "June 24, 1996", "Resolution Date"),
    ("DATE_TIME", "July 4, 1996", "Certificate Date"),
    ("DATE_TIME", "January 13, 2011", "Special Resolution Date"),
    ("DATE_TIME", "August 16, 2011", "RD Order Date"),

    ("ORGANIZATION", "Bhandary Metal Extrusion Private Limited", "Former Company Name"),
    ("ORGANIZATION", "KSH International Private Limited", "Prior Company Name"),
]


def load_document_text(file_path: str) -> str:
    """Loads all paragraphs, table cells, and embedded image OCR text into a single unified text string."""
    doc = docx.Document(file_path)
    text_segments = []

    for p in doc.paragraphs:
        if p.text.strip():
            text_segments.append(p.text.strip())

    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if c.text.strip():
                    text_segments.append(c.text.strip())

    # Extract text from embedded image relationships via EasyOCR
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                img_bytes = rel.target_part.blob
                pil_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                img_np = np.array(pil_img)
                results = ocr_reader.readtext(img_np, detail=0)
                if results:
                    text_segments.append(" ".join(results))
            except Exception:
                pass

    return "\n".join(text_segments)


def run_evaluation(orig_doc_path: str, redacted_doc_path: str):
    """
    Evaluates PII redaction accuracy, precision, recall, and F1-score
    by comparing original text against redacted output.
    """
    print("==================================================")
    print("  EVALUATION BENCHMARK: PII REDACTION SYSTEM")
    print("==================================================")

    orig_text = load_document_text(orig_doc_path)
    redacted_text = load_document_text(redacted_doc_path)

    tp = 0  # True Positives: Sensitive PII correctly redacted
    fn = 0  # False Negatives: Sensitive PII missed in redacted doc
    fp = 0  # False Positives: Non-PII erroneously altered

    category_results: Dict[str, Dict[str, int]] = {}

    for cat, pii_val, context in BENCHMARK_GROUND_TRUTH:
        if cat not in category_results:
            category_results[cat] = {"TP": 0, "FN": 0}

        # Check if the PII string exists in original doc
        in_orig = pii_val in orig_text
        # Check if original PII string still exists in redacted doc
        in_redacted = pii_val in redacted_text

        if in_orig:
            if not in_redacted:
                tp += 1
                category_results[cat]["TP"] += 1
            else:
                fn += 1
                category_results[cat]["FN"] += 1
                print(f"[FN MISSED] Category: {cat:<15} Entity: '{pii_val}'")

    # Evaluate False Positives on non-sensitive domain terms (e.g. 'Companies Act', 'SEBI', 'Offer')
    non_pii_terms = [
        "Companies Act, 2013", "100% Book Built Offer", "Anchor Investors",
        "SEBI ICDR Regulations", "Bid/Offer Closing Day", "Corporate Identity Number"
    ]
    for term in non_pii_terms:
        if term in orig_text and term not in redacted_text:
            fp += 1
            print(f"[FP OVER-REDACTED] Non-PII Term: '{term}'")

    total_eval = tp + fn
    precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
    recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
    f1_score = 2 * (precision * recall) / (precision + recall) if (precision + recall) > 0 else 0.0
    accuracy = tp / (tp + fp + fn) if (tp + fp + fn) > 0 else 0.0

    print("\n--------------------------------------------------")
    print("  CATEGORY-WISE PERFORMANCE BREAKDOWN")
    print("--------------------------------------------------")
    print(f"{'Category':<20} | {'TP':<5} | {'FN':<5} | {'Recall':<8}")
    print("-" * 48)
    for cat, metrics in category_results.items():
        cat_tp = metrics["TP"]
        cat_fn = metrics["FN"]
        cat_rec = cat_tp / (cat_tp + cat_fn) if (cat_tp + cat_fn) > 0 else 0.0
        print(f"{cat:<20} | {cat_tp:<5} | {cat_fn:<5} | {cat_rec:.2%}")

    print("\n--------------------------------------------------")
    print("  OVERALL QUANTITATIVE METRICS")
    print("--------------------------------------------------")
    print(f"  True Positives  (TP) : {tp}")
    print(f"  False Positives (FP) : {fp}")
    print(f"  False Negatives (FN) : {fn}")
    print(f"  ------------------------------")
    print(f"  Accuracy             : {accuracy:.4f} ({accuracy:.2%})")
    print(f"  Precision            : {precision:.4f} ({precision:.2%})")
    print(f"  Recall               : {recall:.4f} ({recall:.2%})")
    print(f"  F1-Score             : {f1_score:.4f} ({f1_score:.2%})")
    print("==================================================\n")

    return {
        "TP": tp, "FP": fp, "FN": fn,
        "Accuracy": accuracy, "Precision": precision,
        "Recall": recall, "F1": f1_score
    }


if __name__ == "__main__":
    orig = "Red Herring Prospectus.docx"
    redacted = "Redacted_Red_Herring_Prospectus.docx"
    run_evaluation(orig, redacted)
