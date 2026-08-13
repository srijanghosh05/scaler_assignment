import io
import os
import re
import sys
import numpy as np
from PIL import Image, ImageDraw
import docx
from docx.opc.constants import RELATIONSHIP_TYPE
import spacy
from faker import Faker

from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_analyzer.nlp_engine import NlpEngineProvider

# Set stdout encoding to UTF-8 for Windows compatibility
sys.stdout.reconfigure(encoding='utf-8')

# Global deterministic PII pseudonymization registry
PII_MAPPING = {}
fake = Faker('en_IN')
Faker.seed(42)  # Reproducible synthetic generation

def get_pseudonym(raw_text: str, entity_type: str) -> str:
    """
    Returns a deterministic synthetic replacement for raw_text based on entity_type.
    Ensures recurring entities always map to the exact same fake identity.
    """
    cleaned_key = raw_text.strip()
    if not cleaned_key:
        return raw_text

    if cleaned_key in PII_MAPPING:
        return PII_MAPPING[cleaned_key]

    # Generate synthetic value based on entity category
    if entity_type in ["PERSON", "PER", "FULL_NAME"]:
        replacement = fake.name()
    elif entity_type in ["EMAIL_ADDRESS", "EMAIL"]:
        replacement = fake.email()
    elif entity_type in ["PHONE_NUMBER", "PHONE"]:
        replacement = "+91 " + fake.numerify("##########")
    elif entity_type == "INDIAN_PAN":
        replacement = fake.bothify(text="?????####?").upper()
    elif entity_type == "INDIAN_AADHAAR":
        replacement = fake.bothify(text="#### #### ####")
    elif entity_type in ["LOCATION", "ADDRESS", "LOC"]:
        replacement = fake.address().replace("\n", ", ")
    elif entity_type in ["ORGANIZATION", "ORG", "COMPANY"]:
        # Only replace if not generic terms
        replacement = fake.company()
    elif entity_type in ["DATE_TIME", "DATE_OF_BIRTH", "DOB"]:
        replacement = fake.date(pattern="%d/%m/%Y")
    elif entity_type in ["CREDIT_CARD", "CREDIT_CARD_NUMBER"]:
        replacement = fake.credit_card_number()
    elif entity_type == "IP_ADDRESS":
        replacement = fake.ipv4()
    else:
        replacement = fake.bothify(text="SYNTHETIC-####")

    PII_MAPPING[cleaned_key] = replacement
    return replacement


def setup_presidio_analyzer() -> AnalyzerEngine:
    """
    Initializes Microsoft Presidio Analyzer with spaCy en_core_web_lg
    and registers custom regex recognizers for Indian domain PII.
    """
    print("[INIT] Loading spaCy en_core_web_lg model into Presidio...")
    configuration = {
        "nlp_engine_name": "spacy",
        "models": [{"lang_code": "en", "model_name": "en_core_web_lg"}]
    }
    provider = NlpEngineProvider(nlp_configuration=configuration)
    nlp_engine = provider.create_engine()
    analyzer = AnalyzerEngine(nlp_engine=nlp_engine, supported_languages=["en"])

    # 1. Indian PAN Recognizer
    pan_pattern = Pattern(
        name="pan_pattern",
        regex=r"\b[A-Z]{5}[0-9]{4}[A-Z]{1}\b",
        score=0.95
    )
    pan_recognizer = PatternRecognizer(
        supported_entity="INDIAN_PAN",
        patterns=[pan_pattern],
        name="Indian_PAN_Recognizer"
    )
    analyzer.registry.add_recognizer(pan_recognizer)

    # 2. Indian Aadhaar Recognizer
    aadhaar_pattern = Pattern(
        name="aadhaar_pattern",
        regex=r"\b[2-9]{1}[0-9]{3}[\s-]?[0-9]{4}[\s-]?[0-9]{4}\b",
        score=0.95
    )
    aadhaar_recognizer = PatternRecognizer(
        supported_entity="INDIAN_AADHAAR",
        patterns=[aadhaar_pattern],
        name="Indian_Aadhaar_Recognizer"
    )
    analyzer.registry.add_recognizer(aadhaar_recognizer)

    # 3. Indian Phone Recognizer
    phone_pattern = Pattern(
        name="phone_pattern",
        regex=r"(\+?91[\s-]?)?[6-9]\d{9}\b|\b0\d{2,4}[\s-]?\d{6,8}\b",
        score=0.90
    )
    phone_recognizer = PatternRecognizer(
        supported_entity="PHONE_NUMBER",
        patterns=[phone_pattern],
        name="Indian_Phone_Recognizer"
    )
    analyzer.registry.add_recognizer(phone_recognizer)

    # 4. Indian / Person Name Recognizer
    person_pattern = Pattern(
        name="person_pattern",
        regex=r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3}\b",
        score=0.55
    )
    person_recognizer = PatternRecognizer(
        supported_entity="PERSON",
        patterns=[person_pattern],
        name="Indian_Person_Recognizer"
    )
    analyzer.registry.add_recognizer(person_recognizer)

    print("[INIT] Presidio Analyzer configured with spaCy + Custom Regex Recognizers.")
    return analyzer


def replace_pii_in_paragraph(paragraph, analyzer: AnalyzerEngine):
    """
    Detects and pseudonymizes PII entity spans within a paragraph object,
    preserving run formatting and avoiding character offset shifting.
    """
    full_text = paragraph.text
    if not full_text or not full_text.strip():
        return

    # Run Presidio analysis
    results = analyzer.analyze(
        text=full_text,
        language="en",
        entities=[
            "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "INDIAN_PAN",
            "INDIAN_AADHAAR", "LOCATION", "ORGANIZATION", "DATE_TIME",
            "CREDIT_CARD", "IP_ADDRESS"
        ],
        score_threshold=0.4
    )

    if not results:
        return

    # Sort matches in reverse order by start index to avoid index shifting
    results = sorted(results, key=lambda r: r.start, reverse=True)

    # Filter out overlapping entity spans
    filtered_results = []
    last_start = len(full_text) + 1
    for r in results:
        if r.end <= last_start:
            filtered_results.append(r)
            last_start = r.start

    # Perform text replacement across runs
    for res in filtered_results:
        raw_val = full_text[res.start:res.end]
        
        # Guard against generic public regulator / metadata headers
        if any(header in raw_val for header in [
            "KSH INTERNATIONAL LIMITED", "SEBI", "RoC", "Companies Act",
            "Registrar of Companies", "Income Tax Department", "100% Book Built Offer",
            "Anchor Investors", "SEBI ICDR Regulations", "Bid/Offer Closing Day",
            "Corporate Identity Number", "Red Herring Prospectus", "SECTION I"
        ]):
            continue

        fake_val = get_pseudonym(raw_val, res.entity_type)

        # Replace text across paragraph runs while preserving styles
        _replace_text_in_runs(paragraph, res.start, res.end, fake_val)

        # Update full_text reference
        full_text = paragraph.text


def _replace_text_in_runs(paragraph, start_idx, end_idx, replacement_text):
    """
    Replaces characters from start_idx to end_idx across paragraph runs
    without destroying run properties (bold, italic, font, etc.).
    """
    curr_pos = 0
    matched_runs = []

    for run in paragraph.runs:
        run_len = len(run.text)
        run_start = curr_pos
        run_end = curr_pos + run_len

        # Check overlap
        if run_end > start_idx and run_start < end_idx:
            matched_runs.append((run, run_start, run_end))

        curr_pos = run_end

    if not matched_runs:
        return

    if len(matched_runs) == 1:
        run, r_start, r_end = matched_runs[0]
        local_start = start_idx - r_start
        local_end = end_idx - r_start
        run.text = run.text[:local_start] + replacement_text + run.text[local_end:]
    else:
        # Spans across multiple runs: place replacement in first run, clear overlapping middle/end runs
        first_run, f_start, _ = matched_runs[0]
        local_start = start_idx - f_start
        first_run.text = first_run.text[:local_start] + replacement_text

        for run, r_start, r_end in matched_runs[1:-1]:
            run.text = ""

        last_run, l_start, _ = matched_runs[-1]
        local_end = end_idx - l_start
        last_run.text = last_run.text[local_end:]


def redact_docx_tables(doc, analyzer: AnalyzerEngine):
    """
    Iterates recursively over all document tables and table cells to redact PII.
    """
    print(f"[PROCESS] Scanning {len(doc.tables)} document tables...")
    cell_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_count += 1
                for paragraph in cell.paragraphs:
                    replace_pii_in_paragraph(paragraph, analyzer)
    print(f"[PROCESS] Processed {cell_count} table cells successfully.")


def redact_docx_images(doc, analyzer: AnalyzerEngine):
    """
    Extracts embedded image parts from document relationships, performs EasyOCR,
    and draws solid redaction bounding boxes over detected PII coordinates.
    """
    print("[PROCESS] Initializing EasyOCR for embedded image inspection...")
    import easyocr
    reader = easyocr.Reader(['en'], gpu=False, verbose=False)

    image_parts_processed = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            img_part = rel.target_part
            img_bytes = img_part.blob
            
            try:
                pil_img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                img_np = np.array(pil_img)
                
                # Perform OCR detection
                ocr_results = reader.readtext(img_np)
                
                if not ocr_results:
                    continue

                draw = ImageDraw.Draw(pil_img)
                redacted_any = False

                for bbox, text, prob in ocr_results:
                    if prob < 0.2 or not text.strip():
                        continue

                    # Analyze detected text string for PII
                    analysis = analyzer.analyze(
                        text=text,
                        language="en",
                        entities=[
                            "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "INDIAN_PAN",
                            "INDIAN_AADHAAR", "LOCATION", "ORGANIZATION", "DATE_TIME"
                        ],
                        score_threshold=0.4
                    )

                    # Also check explicit regex triggers for card text (Names, PAN, Aadhaar, DOB, Addresses)
                    is_pii = len(analysis) > 0 or any([
                        re.search(r"[A-Z]{5}[0-9]{4}[A-Z]{1}", text),  # PAN
                        re.search(r"\d{4}\s?\d{4}\s?\d{4}", text),     # Aadhaar
                        re.search(r"\d{2}/\d{2}/\d{4}", text),         # DOB
                        any(kw in text.upper() for kw in ["FATHER", "NAME", "DOB", "ADDRESS", "PERMANENT"])
                    ])

                    if is_pii:
                        redacted_any = True
                        # Get bounding box points [(x1,y1), (x2,y2), (x3,y3), (x4,y4)]
                        x_coords = [p[0] for p in bbox]
                        y_coords = [p[1] for p in bbox]
                        x_min, x_max = int(min(x_coords)), int(max(x_coords))
                        y_min, y_max = int(min(y_coords)), int(max(y_coords))

                        # Draw solid black redaction box with padding
                        pad = 3
                        draw.rectangle(
                            [x_min - pad, y_min - pad, x_max + pad, y_max + pad],
                            fill="black"
                        )

                if redacted_any:
                    out_buffer = io.BytesIO()
                    pil_img.save(out_buffer, format="PNG")
                    img_part._blob = out_buffer.getvalue()
                    image_parts_processed += 1
                    print(f"  -> Redacted PII in embedded image part: {rel.target_ref}")

            except Exception as e:
                print(f"  [WARN] Could not process image {rel.target_ref}: {e}")

    print(f"[PROCESS] Embedded image redaction completed ({image_parts_processed} images updated).")


def process_docx_redaction(input_path: str, output_path: str):
    """
    Main execution pipeline for docx PII redaction.
    """
    print(f"\n==================================================")
    print(f"  STARTING PII REDACTION PIPELINE")
    print(f"  Input Document : {input_path}")
    print(f"  Output Document: {output_path}")
    print(f"==================================================\n")

    analyzer = setup_presidio_analyzer()
    doc = docx.Document(input_path)

    # 1. Redact standard paragraphs
    print(f"[PROCESS] Scanning {len(doc.paragraphs)} document paragraphs...")
    for idx, p in enumerate(doc.paragraphs):
        replace_pii_in_paragraph(p, analyzer)

    # 2. Redact table cells
    redact_docx_tables(doc, analyzer)

    # 3. Redact embedded images via OCR
    redact_docx_images(doc, analyzer)

    # Save output document
    doc.save(output_path)
    print(f"\n[SUCCESS] Document PII redaction complete!")
    print(f"[SUCCESS] Redacted file saved as: '{output_path}'")
    print(f"[INFO] Total Unique PII Pseudonymized: {len(PII_MAPPING)} entities.")


if __name__ == "__main__":
    input_doc = "Red Herring Prospectus.docx"
    output_doc = "Redacted_Red_Herring_Prospectus.docx"
    process_docx_redaction(input_doc, output_doc)
