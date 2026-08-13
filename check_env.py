import sys
import docx
from presidio_analyzer import AnalyzerEngine, PatternRecognizer, Pattern
from presidio_anonymizer import AnonymizerEngine
import spacy
from faker import Faker
import easyocr
import PIL

# Ensure UTF-8 output encoding for Windows terminal unicode characters
sys.stdout.reconfigure(encoding='utf-8')

def main():
    print("=== PII Redaction Environment Verification ===")
    print(f"Python Version: {sys.version.split()[0]}")

    # 1. Test docx reading capability
    doc_path = "Enterprise Data - Assignment.docx"
    try:
        doc = docx.Document(doc_path)
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        print(f"[OK] docx module loaded successfully. Read '{doc_path}' ({paragraph_count} paragraphs, {table_count} tables).")
    except Exception as e:
        print(f"[ERROR] docx verification failed: {e}")

    # 2. Test Presidio Analyzer & Custom Regex
    try:
        analyzer = AnalyzerEngine()
        
        # Test custom regex recognizer (e.g. Indian PAN / Aadhaar pattern test)
        pan_pattern = Pattern(name="pan_pattern", regex=r"[A-Z]{5}[0-9]{4}[A-Z]{1}", score=0.85)
        pan_recognizer = PatternRecognizer(supported_entity="INDIAN_PAN", patterns=[pan_pattern])
        analyzer.registry.add_recognizer(pan_recognizer)
        
        results = analyzer.analyze(text="My PAN is ABCDE1234F", language="en")
        entities_found = [r.entity_type for r in results]
        print(f"[OK] presidio_analyzer initialized with custom regex. Detected entities: {entities_found}")
    except Exception as e:
        print(f"[ERROR] Presidio verification failed: {e}")

    # 3. Test spaCy models loading
    for model_name in ["en_core_web_sm", "en_core_web_lg"]:
        try:
            nlp = spacy.load(model_name)
            doc_spacy = nlp("Google Antigravity PII Redaction Engine")
            ents = [(ent.text, ent.label_) for ent in doc_spacy.ents]
            print(f"[OK] spacy loaded model '{model_name}' successfully. NER test: {ents}")
        except Exception as e:
            print(f"[ERROR] spacy model '{model_name}' failed to load: {e}")

    # 4. Test Faker with 'en_IN' locale
    try:
        fake = Faker('en_IN')
        sample_name = fake.name()
        sample_aadhaar = fake.bothify(text='#### #### ####')
        print(f"[OK] faker ('en_IN') loaded. Sample data -> Name: {sample_name}, Aadhaar placeholder: {sample_aadhaar}")
    except Exception as e:
        print(f"[ERROR] faker verification failed: {e}")

    # 5. Test EasyOCR Engine
    try:
        reader = easyocr.Reader(['en'], gpu=False, verbose=False)
        print(f"[OK] easyocr initialized successfully.")
    except Exception as e:
        print(f"[ERROR] easyocr verification failed: {e}")

if __name__ == "__main__":
    main()
